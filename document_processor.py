"""
Document Processing Module for RAG Chatbot
Handles file uploads, text extraction, and document chunking
"""

import streamlit as st
import tempfile
import os
from typing import List, Dict, Any
from io import BytesIO

# Document processing imports
import PyPDF2
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document as LangChainDocument


class DocumentProcessor:
    """Handles document processing for RAG implementation"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize the document processor
        
        Args:
            chunk_size: Size of text chunks for processing
            chunk_overlap: Overlap between chunks to maintain context
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def extract_text_from_pdf(self, file_bytes: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_file = BytesIO(file_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
            
            return text.strip()
        except Exception as e:
            st.error(f"Error extracting text from PDF: {str(e)}")
            return ""
    
    def extract_text_from_docx(self, file_bytes: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            docx_file = BytesIO(file_bytes)
            doc = Document(docx_file)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text.strip()
        except Exception as e:
            st.error(f"Error extracting text from DOCX: {str(e)}")
            return ""
    
    def extract_text_from_txt(self, file_bytes: bytes) -> str:
        """Extract text from TXT file"""
        try:
            text = file_bytes.decode('utf-8')
            return text.strip()
        except UnicodeDecodeError:
            try:
                text = file_bytes.decode('latin-1')
                return text.strip()
            except Exception as e:
                st.error(f"Error extracting text from TXT: {str(e)}")
                return ""
        except Exception as e:
            st.error(f"Error extracting text from TXT: {str(e)}")
            return ""
    
    def extract_text(self, uploaded_file) -> str:
        """
        Extract text from uploaded file based on file type
        
        Args:
            uploaded_file: Streamlit uploaded file object
        
        Returns:
            Extracted text as string
        """
        file_type = uploaded_file.type
        file_bytes = uploaded_file.read()
        
        if file_type == "application/pdf":
            return self.extract_text_from_pdf(file_bytes)
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return self.extract_text_from_docx(file_bytes)
        elif file_type == "text/plain":
            return self.extract_text_from_txt(file_bytes)
        else:
            st.error(f"Unsupported file type: {file_type}")
            return ""
    
    def chunk_text(self, text: str, metadata: Dict[str, Any] = None) -> List[LangChainDocument]:
        """
        Split text into chunks for processing
        
        Args:
            text: Text to be chunked
            metadata: Additional metadata for the document
        
        Returns:
            List of LangChain Document objects
        """
        if not text.strip():
            return []
        
        # Split text into chunks
        chunks = self.text_splitter.split_text(text)
        
        # Create LangChain documents with metadata
        documents = []
        for i, chunk in enumerate(chunks):
            doc_metadata = metadata.copy() if metadata else {}
            doc_metadata.update({
                "chunk_index": i,
                "chunk_size": len(chunk)
            })
            
            documents.append(LangChainDocument(
                page_content=chunk,
                metadata=doc_metadata
            ))
        
        return documents
    
    def process_uploaded_files(self, uploaded_files: List) -> List[LangChainDocument]:
        """
        Process multiple uploaded files and return document chunks
        
        Args:
            uploaded_files: List of Streamlit uploaded file objects
        
        Returns:
            List of processed document chunks
        """
        all_documents = []
        
        for uploaded_file in uploaded_files:
            if uploaded_file is not None:
                # Extract text from file
                text = self.extract_text(uploaded_file)
                
                if text:
                    # Create metadata for the document
                    metadata = {
                        "source": uploaded_file.name,
                        "file_type": uploaded_file.type,
                        "file_size": len(uploaded_file.getvalue())
                    }
                    
                    # Chunk the text
                    documents = self.chunk_text(text, metadata)
                    all_documents.extend(documents)
                    
                    st.success(f"✅ Processed {uploaded_file.name}: {len(documents)} chunks created")
                else:
                    st.warning(f"⚠️ No text extracted from {uploaded_file.name}")
        
        return all_documents
    
    def get_supported_file_types(self) -> List[str]:
        """Return list of supported file types"""
        return ["pdf", "docx", "txt"]
    
    def validate_file(self, uploaded_file) -> bool:
        """
        Validate if the uploaded file is supported
        
        Args:
            uploaded_file: Streamlit uploaded file object
        
        Returns:
            True if file is supported, False otherwise
        """
        if uploaded_file is None:
            return False
        
        supported_types = {
            "application/pdf": "pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
            "text/plain": "txt"
        }
        
        return uploaded_file.type in supported_types


def create_document_processor() -> DocumentProcessor:
    """Factory function to create a DocumentProcessor instance"""
    return DocumentProcessor()


# Example usage and testing function
def test_document_processor():
    """Test function for the DocumentProcessor class"""
    processor = create_document_processor()
    
    st.write("### Document Processor Test")
    st.write(f"Supported file types: {', '.join(processor.get_supported_file_types())}")
    st.write(f"Chunk size: {processor.chunk_size}")
    st.write(f"Chunk overlap: {processor.chunk_overlap}")


if __name__ == "__main__":
    # This section runs when the file is executed directly
    test_document_processor()